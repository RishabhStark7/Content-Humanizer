"""Exporter for publication-ready Word (.docx) documents."""

from pathlib import Path
from typing import Union
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from schemas.document import DocumentModel
from schemas.validation import ValidationReport


def export_to_docx(
    document: DocumentModel,
    output_path: Union[str, Path],
    validation_report: ValidationReport = None
) -> Path:
    """Export synthesized DocumentModel to a publication-ready .docx document.

    Args:
        document: Synthesized DocumentModel.
        output_path: Target file path (.docx).
        validation_report: Optional ValidationReport to include metadata header.

    Returns:
        Path of saved .docx file.
    """
    out_file = Path(output_path)
    out_file.parent.mkdir(parents=True, exist_ok=True)

    doc = docx.Document()

    # Document Title
    h1 = doc.add_heading(document.title, level=0)
    h1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h1.runs:
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)  # Dark Gray

    # Subtitle / Metadata callout
    meta_p = doc.add_paragraph()
    meta_run = meta_p.add_run("Synthesized by Human Writing Engine | Publication-Ready Standard")
    meta_run.font.italic = True
    meta_run.font.size = Pt(9.5)
    meta_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()  # Spacer

    # Render Sections
    for sec in document.sections:
        h = doc.add_heading(sec.heading, level=min(sec.level, 3))
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

        p = doc.add_paragraph(sec.content)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(8)

        # Bullets
        for b in sec.bullets:
            bp = doc.add_paragraph(b, style="List Bullet")
            bp.paragraph_format.space_after = Pt(4)

    # Render FAQs if present
    if document.faqs:
        faq_h = doc.add_heading("Frequently Asked Questions (FAQs)", level=2)
        for run in faq_h.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

        for faq in document.faqs:
            qh = doc.add_heading(faq.question, level=3)
            for run in qh.runs:
                run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

            ans_p = doc.add_paragraph(faq.answer)
            ans_p.paragraph_format.line_spacing = 1.15
            ans_p.paragraph_format.space_after = Pt(8)

    try:
        doc.save(str(out_file))
        return out_file
    except PermissionError:
        alt_file = out_file.parent / f"{out_file.stem}_new.docx"
        doc.save(str(alt_file))
        return alt_file
